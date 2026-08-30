"""

匯出檔案產生器：把分類結果的 rows（大類別、子類別、問卷回覆內容、判斷原因與
說明、受試者建議摘要）產生成真正的 Excel（.xlsx）或 Word（.docx）檔案。

CSV 沒有放在這裡——CSV 是純文字格式，前端組字串就能產生，不需要動用
後端套件；這裡只處理 Excel、Word 這兩種需要真正函式庫才能產生的二進位
格式。

兩種格式都遵循跟 CSV 匯出、前端表格顯示一致的規則：
  - 5 欄：大類別、子類別、問卷回覆內容、判斷原因與說明、受試者建議摘要
  - 大類別連續相同時合併儲存格（Excel 用 merge_cells，Word 用手動合併
    儲存格），不逐列重複
"""

import io

import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT

COLUMN_HEADERS = ["大類別", "子類別", "問卷回覆內容", "判斷原因與說明", "受試者建議摘要"]


def _row_values(row: dict) -> list:
    return [
        row.get("main_category", ""),
        row.get("sub_category", ""),
        row.get("respondent_text", ""),
        row.get("aggregated_reasoning", ""),
        row.get("aggregated_summary", ""),
    ]


def build_xlsx(rows: list, title: str = "分類結果") -> bytes:
    """把 rows 產生成 .xlsx 檔案，回傳檔案的原始 bytes。"""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = title[:31] if title else "分類結果"  # Excel 分頁名稱上限 31 字元

    header_font = Font(name="微軟正黑體", bold=True, color="FFFFFF")
    header_fill = PatternFill(start_color="F43F5E", end_color="F43F5E", fill_type="solid")
    body_font = Font(name="微軟正黑體")
    wrap_alignment = Alignment(wrap_text=True, vertical="top", horizontal="left")
    center_alignment = Alignment(wrap_text=True, vertical="center", horizontal="center")

    for col_idx, header in enumerate(COLUMN_HEADERS, start=1):
        cell = ws.cell(row=1, column=col_idx, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = center_alignment

    for row_idx, row in enumerate(rows, start=2):
        values = _row_values(row)
        for col_idx, value in enumerate(values, start=1):
            cell = ws.cell(row=row_idx, column=col_idx, value=value)
            cell.font = body_font
            cell.alignment = wrap_alignment

    # 大類別（第一欄）連續相同時合併儲存格，比照前端畫面上的 rowSpan 效果
    merge_start_row = 2
    for row_idx in range(3, len(rows) + 3):
        current_main = rows[row_idx - 2]["main_category"] if row_idx - 2 < len(rows) else None
        prev_main = rows[row_idx - 3]["main_category"] if row_idx - 3 < len(rows) else None
        if current_main != prev_main:
            if row_idx - 1 > merge_start_row:
                ws.merge_cells(start_row=merge_start_row, start_column=1, end_row=row_idx - 1, end_column=1)
                ws.cell(row=merge_start_row, column=1).alignment = center_alignment
            merge_start_row = row_idx

    column_widths = [16, 20, 46, 40, 34]
    for col_idx, width in enumerate(column_widths, start=1):
        ws.column_dimensions[get_column_letter(col_idx)].width = width

    ws.freeze_panes = "A2"

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def build_docx(rows: list, title: str = "分類結果") -> bytes:
    """把 rows 產生成 .docx 檔案，回傳檔案的原始 bytes。"""
    doc = Document()

    # 頁面設定：橫向，欄位比較多，直向會太窄
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    heading = doc.add_heading(title or "分類結果", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    for i, header in enumerate(COLUMN_HEADERS):
        hdr_cells[i].text = header
        for p in hdr_cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(11)

    # 欄寬（單位 cm，5 欄加總跟橫向可用寬度大致相符）
    col_widths_cm = [2.2, 2.8, 7.5, 7.0, 6.0]

    row_start_index = {}  # main_category -> 這個 docx table 裡的起始 row index
    prev_main_category = None
    for r_idx, row in enumerate(rows):
        cells = table.add_row().cells
        values = _row_values(row)
        main_category = row.get("main_category", "")
        is_continuation = r_idx > 0 and prev_main_category == main_category

        for c_idx, value in enumerate(values):
            # 【修正】大類別（第 0 欄）如果這列要被合併進上一列，這裡故意
            # 先不寫文字——python-docx 合併兩個「都已經有文字」的儲存格時，
            # 會把兩段文字直接接在一起（變成「主管領導\n主管領導」），
            # 不是保留第一格、丟棄第二格。留空再合併，才不會重複。
            if c_idx == 0 and is_continuation:
                continue
            cells[c_idx].text = str(value)
            for p in cells[c_idx].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
            cells[c_idx].vertical_alignment = WD_ALIGN_VERTICAL.TOP

        table_row_idx = r_idx + 1  # +1 因為第 0 列是表頭
        if is_continuation:
            first_row_idx = row_start_index[main_category]
            first_cell = table.rows[first_row_idx].cells[0]
            this_cell = table.rows[table_row_idx].cells[0]
            first_cell.merge(this_cell)
            for p in first_cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            row_start_index[main_category] = table_row_idx
            cells[0].vertical_alignment = WD_ALIGN_VERTICAL.TOP
            for p in cells[0].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(10)

        prev_main_category = main_category

    for row in table.rows:
        for i, cm in enumerate(col_widths_cm):
            row.cells[i].width = Cm(cm)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()